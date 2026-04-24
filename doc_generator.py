"""
문서 생성 모듈
python-docx를 사용하여 신한 금융시장 Brief 양식의
Word(.docx) 문서를 생성합니다.

두 가지 모드 지원:
  1. generate()              - 처음부터 프로그래매틱하게 문서 생성
  2. generate_from_template() - .doc/.docx 양식 파일에 내용 채워넣기
"""

import copy
import datetime
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class DocGenerator:
    """신한 금융시장 Brief 양식 Word 문서 생성기"""

    # 색상 상수
    COLOR_BLACK = RGBColor(0, 0, 0)
    COLOR_DARK_GRAY = RGBColor(80, 80, 80)
    COLOR_GRAY = RGBColor(128, 128, 128)
    COLOR_DARK_BLUE = RGBColor(0, 51, 102)

    def __init__(self, config: dict):
        self.config = config
        self.issue = config["issue"]

    # ─── 유틸리티 메서드 ───

    def _get_date_info(self) -> tuple[str, str]:
        """날짜 문자열과 파일명용 날짜 접두사 반환"""
        date_config = self.issue["date"]
        if date_config == "auto":
            today = datetime.date.today()
        else:
            today = datetime.datetime.strptime(date_config, "%Y-%m-%d").date()

        # "2026년 2월 7일" 형식
        date_str = f"{today.year}년 {today.month}월 {today.day}일"
        # "260207" 형식 (파일명용)
        date_prefix = today.strftime("%y%m%d")
        return date_str, date_prefix

    @staticmethod
    def _set_font(run, font_name: str = "맑은 고딕", size_pt: float = 10,
                  bold: bool = False, color: RGBColor = None):
        """run에 한글 호환 폰트 설정"""
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        if color:
            run.font.color.rgb = color

        # 동아시아(한글) 폰트 설정
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _set_paragraph_spacing(paragraph, before_pt: float = 0,
                               after_pt: float = 0, line_pt: float = None):
        """문단 간격 설정"""
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(before_pt)
        fmt.space_after = Pt(after_pt)
        if line_pt is not None:
            fmt.line_spacing = Pt(line_pt)

    @staticmethod
    def _add_bottom_border(paragraph, color: str = "333333", size: str = "6"):
        """문단 아래에 수평선 추가"""
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): size,
            qn("w:space"): "1",
            qn("w:color"): color,
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ─── 문서 섹션 생성 메서드 ───

    def _add_header(self, doc: Document):
        """문서 상단: '사내한 / 대외비' 라벨"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_paragraph_spacing(p, before_pt=0, after_pt=2)
        run = p.add_run(self.issue["confidential"])
        self._set_font(run, size_pt=8, color=self.COLOR_GRAY)

    def _add_title_block(self, doc: Document, date_str: str, issue_num: int):
        """메인 타이틀 블록: 제목 + 부제 + 호수/날짜"""
        # 메인 타이틀
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._set_paragraph_spacing(p, before_pt=0, after_pt=0)

        run = p.add_run(self.issue["title"])
        self._set_font(run, size_pt=18, bold=True, color=self.COLOR_DARK_BLUE)

        run = p.add_run(f"  {self.issue['org']}")
        self._set_font(run, size_pt=9, color=self.COLOR_GRAY)

        # 영문 부제 + 호수/날짜
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._set_paragraph_spacing(p, before_pt=0, after_pt=4)

        run = p.add_run(
            f"{self.issue['subtitle']} {self.issue['label']}"
        )
        self._set_font(run, size_pt=8.5, color=self.COLOR_GRAY)
        run.italic = True

        run = p.add_run(f"  {issue_num}호 / {date_str}")
        self._set_font(run, size_pt=8.5, color=self.COLOR_DARK_GRAY)

        # 구분선
        separator = doc.add_paragraph()
        self._set_paragraph_spacing(separator, before_pt=0, after_pt=6)
        self._add_bottom_border(separator)

    def _add_article(self, doc: Document, article: dict):
        """개별 기사 블록: 제목(볼드) + 불릿(●) 요약 3줄"""
        title = article.get("title", "")
        lines = article.get("lines", [])

        # 기사 제목
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._set_paragraph_spacing(p, before_pt=8, after_pt=1)
        run = p.add_run(title)
        self._set_font(run, size_pt=10, bold=True, color=self.COLOR_BLACK)

        # 요약 줄 (불릿 포인트)
        for line in lines:
            if not line or not line.strip():
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=14)
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run(f"● {line}")
            self._set_font(run, size_pt=9, color=self.COLOR_BLACK)

    # ─── 템플릿 기반 생성 메서드 ───

    @staticmethod
    def _convert_doc_to_docx(doc_path: str) -> str:
        """macOS textutil을 사용하여 .doc -> .docx 변환"""
        doc_path = Path(doc_path)
        docx_path = doc_path.with_suffix(".docx")
        # 이미 .docx로 변환된 파일이 있으면 재사용
        if not docx_path.exists() or docx_path.stat().st_mtime < doc_path.stat().st_mtime:
            subprocess.run(
                ["textutil", "-convert", "docx", str(doc_path), "-output", str(docx_path)],
                check=True, capture_output=True,
            )
        return str(docx_path)

    def _replace_paragraph_text(self, paragraph, new_text: str, 
                                font_name: str = None):
        """
        문단 텍스트를 교체하고 폰트 적용.
        
        Args:
            paragraph: 교체할 문단
            new_text: 새 텍스트
            font_name: 적용할 폰트명 (None이면 원본 유지)
        """
        if not paragraph.runs:
            run = paragraph.add_run(new_text)
            if font_name:
                self._set_font(run, font_name=font_name)
            return
        
        # 첫 번째 run의 서식 정보 백업
        first_run = paragraph.runs[0]
        orig_font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
        font_color = first_run.font.color.rgb if first_run.font.color.rgb else None
        
        # 텍스트 교체
        first_run.text = new_text
        
        # 폰트 적용 (지정된 폰트 또는 원본)
        target_font = font_name if font_name else orig_font_name
        if target_font:
            first_run.font.name = target_font
            # 한글 폰트 설정
            rPr = first_run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rPr.makeelement(qn("w:rFonts"), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), target_font)
        
        # 서식 복원
        if font_size:
            first_run.font.size = font_size
        if font_bold is not None:
            first_run.font.bold = font_bold
        if font_italic is not None:
            first_run.font.italic = font_italic
        if font_color:
            first_run.font.color.rgb = font_color
        
        # 나머지 run 제거
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)

    def _find_table_cells_with_placeholders(self, doc: Document) -> list[dict]:
        """
        테이블 기반 양식에서 {{T}} 번호별로 기사 슬롯을 찾음.
        같은 {{T}} 번호를 가진 병합 셀은 하나의 슬롯으로 취급.
        반환: [{table_idx, row_idx, col_idx, cell, paragraphs, t_number}]
        """
        import re
        
        # {{T}} 번호별로 첫 번째 셀만 수집
        t_number_to_cell = {}  # {t_number: cell_info}
        
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    
                    # {{T숫자}} 추출
                    match = re.search(r'\{\{T(\d+)\}\}', text)
                    if match:
                        t_num = int(match.group(1))
                        
                        # 이 {{T}} 번호가 처음 발견되면 저장
                        if t_num not in t_number_to_cell:
                            t_number_to_cell[t_num] = {
                                "table_idx": t_idx,
                                "row_idx": r_idx,
                                "col_idx": c_idx,
                                "cell": cell,
                                "paragraphs": cell.paragraphs,
                                "t_number": t_num,
                            }
        
        # {{T}} 번호 순서대로 정렬하여 반환
        sorted_cells = [t_number_to_cell[num] for num in sorted(t_number_to_cell.keys())]
        
        return sorted_cells
    
    def _parse_cell_structure(self, cell_paragraphs: list) -> dict:
        """
        셀 내부 문단 리스트를 분석하여 제목/불릿/저자 인덱스 반환.
        {{T}}, {{B}}, ~~ 조합을 유연하게 처리.
        반환: {title_para_idx, bullet_para_idxs: [idx1, idx2, idx3], author_idx}
        """
        title_idx = None
        bullet_idxs = []
        tilde_indices = []
        author_idx = None
        
        # 1단계: {{T}}, {{B}}, ~~, (XXX Tel) 인덱스 수집
        for i, para in enumerate(cell_paragraphs):
            text = para.text.strip()
            
            if text.startswith("{{T"):
                title_idx = i
            elif text.startswith("{{B"):
                bullet_idxs.append(i)
            elif text == "~~":
                tilde_indices.append(i)
            elif "Tel" in text and "(" in text:
                # (XXX Tel xxxx) 형태의 저자 문단
                author_idx = i
        
        # 2단계: {{T}}가 있으면 제목으로 사용하고, {{B}} + ~~ 를 합쳐 불릿 3개 만들기
        if title_idx is not None:
            # {{B}}와 ~~ 중에서 제목 뒤에 나오는 것들을 불릿으로 사용
            all_bullets = []
            for idx in bullet_idxs:
                if idx > title_idx:
                    all_bullets.append(idx)
            for idx in tilde_indices:
                if idx > title_idx:
                    all_bullets.append(idx)
            all_bullets.sort()  # 순서대로 정렬
            return {
                "title_idx": title_idx,
                "bullet_idxs": all_bullets[:3],
                "author_idx": author_idx
            }
        
        # 3단계: {{T}}가 없고 ~~만 있으면 첫 4개 ~~를 제목+불릿으로 사용
        if len(tilde_indices) >= 4:
            return {
                "title_idx": tilde_indices[0],
                "bullet_idxs": tilde_indices[1:4],
                "author_idx": author_idx
            }
        
        # 4단계: 위의 경우에 해당하지 않으면 빈 구조 반환
        return {
            "title_idx": title_idx,
            "bullet_idxs": bullet_idxs[:3],
            "author_idx": author_idx
        }
    
    def _replace_cell_paragraph_text(self, paragraph, new_text: str, 
                                     font_name: str = None,
                                     font_size_pt: float = None,
                                     font_bold: bool = None):
        """
        셀 내부 문단의 텍스트를 교체하고 폰트 적용.
        
        Args:
            paragraph: 교체할 문단
            new_text: 새 텍스트
            font_name: 적용할 폰트명 (None이면 원본 유지)
            font_size_pt: 폰트 크기 (포인트, None이면 원본 유지)
            font_bold: 볼드 여부 (None이면 원본 유지)
        """
        # 첫 번째 run의 서식 백업
        orig_font_name = None
        orig_font_size = None
        orig_font_bold = None
        font_italic = None
        font_color = None
        
        if paragraph.runs:
            first_run = paragraph.runs[0]
            orig_font_name = first_run.font.name
            orig_font_size = first_run.font.size
            orig_font_bold = first_run.font.bold
            font_italic = first_run.font.italic
            font_color = first_run.font.color.rgb if first_run.font.color.rgb else None
        
        # 기존 텍스트 모두 제거
        for run in paragraph.runs:
            run.text = ""
        
        # 새 텍스트 추가 (첫 run 재활용 또는 새로 생성)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.text = new_text
        else:
            first_run = paragraph.add_run(new_text)
        
        # 폰트명 적용
        target_font = font_name if font_name else orig_font_name
        if target_font:
            first_run.font.name = target_font
            # 한글 폰트 설정
            rPr = first_run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rPr.makeelement(qn("w:rFonts"), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), target_font)
        
        # 폰트 크기 적용
        target_size = Pt(font_size_pt) if font_size_pt else orig_font_size
        if target_size:
            first_run.font.size = target_size
        
        # 볼드 적용
        target_bold = font_bold if font_bold is not None else orig_font_bold
        if target_bold is not None:
            first_run.font.bold = target_bold
        
        # 기타 서식 복원
        if font_italic is not None:
            first_run.font.italic = font_italic
        if font_color:
            first_run.font.color.rgb = font_color

    def _clone_article_group(self, doc: Document, ref_slot: dict, insert_after_idx: int) -> dict:
        """
        기존 기사 슬롯의 XML을 복제하여 insert_after_idx 뒤에 삽입.
        새로 생성된 슬롯 정보를 반환.
        """
        body = doc.element.body
        paras = doc.paragraphs

        # 복제 대상: 제목 ~ 저자까지
        start = ref_slot["title_idx"]
        end = ref_slot["author_idx"]
        ref_elements = []
        for idx in range(start, end + 1):
            ref_elements.append(paras[idx]._element)

        # 삽입 위치: insert_after_idx 문단 바로 뒤
        anchor = paras[insert_after_idx]._element

        new_slot = {"title_idx": None, "bullet_idxs": [], "author_idx": None}
        for j, elem in enumerate(ref_elements):
            new_elem = copy.deepcopy(elem)
            anchor.addnext(new_elem)
            anchor = new_elem  # 다음 요소는 이 뒤에 삽입

        return new_slot  # 인덱스는 나중에 재계산

    def generate_from_template(self, articles: list[dict],
                               template_path: str = "금융시장브리프 양식.doc",
                               output_path: str = None) -> str:
        """
        테이블 기반 양식 파일을 열어 기사 내용을 채워넣은 뒤 저장.

        Args:
            articles: 기사 목록
            template_path: 양식 파일 경로 (.doc 또는 .docx)
            output_path: 저장 경로 (None이면 자동 생성)

        Returns:
            저장된 파일 경로
        """
        date_str, date_prefix = self._get_date_info()
        issue_num = self.issue["number"]

        # .doc이면 .docx로 변환
        tpl_path = Path(template_path)
        if tpl_path.suffix.lower() == ".doc":
            docx_tpl = self._convert_doc_to_docx(template_path)
        else:
            docx_tpl = str(tpl_path)

        doc = Document(docx_tpl)

        # ── 1) 플레이스홀더 셀 찾기 (문서 변경 전에 먼저 실행) ──
        content_cells = self._find_table_cells_with_placeholders(doc)
        
        print(f"  양식에서 {len(content_cells)}개 기사 슬롯 발견")
        
        if len(content_cells) < len(articles):
            print(f"  경고: 양식 슬롯({len(content_cells)}개)보다 기사({len(articles)}개)가 많습니다.")
            print(f"        처음 {len(content_cells)}개 기사만 채워집니다.")

        # ── 2) 호수/날짜 업데이트 ──
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "호 /" in para.text and "년" in para.text:
                            self._replace_cell_paragraph_text(
                                para, f"{issue_num}호 / {date_str}"
                            )

        # ── 3) 각 셀에 기사 내용 채우기 ──
        for i, cell_info in enumerate(content_cells):
            if i >= len(articles):
                break
            
            article = articles[i]
            title = article.get("title", "")
            lines = article.get("lines", [])
            author = article.get("author", "")
            
            # 셀 내부 구조 파싱
            structure = self._parse_cell_structure(cell_info["paragraphs"])
            
            # 제목 채우기 (나눔고딕 볼드체, 크기 12)
            if structure["title_idx"] is not None:
                title_para = cell_info["paragraphs"][structure["title_idx"]]
                self._replace_cell_paragraph_text(
                    title_para, title, 
                    font_name="나눔고딕", 
                    font_size_pt=12, 
                    font_bold=True
                )
            
            # 불릿 3줄 채우기 (HY신명조, 크기 10.5, ● 기호 제거)
            for j, bullet_idx in enumerate(structure["bullet_idxs"]):
                if j < len(lines) and lines[j]:
                    bullet_para = cell_info["paragraphs"][bullet_idx]
                    self._replace_cell_paragraph_text(
                        bullet_para, lines[j], 
                        font_name="HY신명조", 
                        font_size_pt=10.5
                    )
            
            # 저자 정보 채우기 (나눔바른고딕, 크기 9)
            if structure["author_idx"] is not None and author:
                author_para = cell_info["paragraphs"][structure["author_idx"]]
                self._replace_cell_paragraph_text(
                    author_para, f"({author})", 
                    font_name="나눔바른고딕", 
                    font_size_pt=9
                )

        # ── 4) 저장 ──
        if output_path is None:
            output_path = (
                f"{date_prefix}_Shinhan Financial Market Brief"
                f"_{issue_num}호.docx"
            )

        doc.save(output_path)
        return output_path

    # ─── 프로그래매틱 생성 메서드 ───

    def generate(self, articles: list[dict], output_path: str = None) -> str:
        """
        전체 문서를 생성하고 저장합니다.

        Args:
            articles: 기사 목록 (각 항목에 title, lines, source 포함)
            output_path: 저장 경로 (None이면 자동 생성)

        Returns:
            저장된 파일 경로
        """
        doc = Document()
        date_str, date_prefix = self._get_date_info()
        issue_num = self.issue["number"]

        # ── 페이지 설정 (A4) ──
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # ── 기본 폰트 스타일 설정 ──
        style = doc.styles["Normal"]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(10)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "맑은 고딕")

        # ── 문서 구조 생성 ──
        self._add_header(doc)
        self._add_title_block(doc, date_str, issue_num)

        for article in articles:
            self._add_article(doc, article)

        # ── 저장 ──
        if output_path is None:
            output_path = (
                f"{date_prefix}_Shinhan Financial Market Brief"
                f"_{issue_num}호.docx"
            )

        doc.save(output_path)
        return output_path
